from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

COLOR_TITLE = RGBColor(0x07, 0x44, 0x34)
COLOR_HIGH  = RGBColor(0xDC, 0x26, 0x26)
COLOR_MED   = RGBColor(0xD9, 0x77, 0x06)
COLOR_LOW   = RGBColor(0x16, 0xA3, 0x4A)
COLOR_GRAY  = RGBColor(0x6B, 0x72, 0x80)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_DARK  = RGBColor(0x11, 0x18, 0x27)
COLOR_BODY  = RGBColor(0x1F, 0x29, 0x37)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_TITLE


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_BODY


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_BODY


def issue_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_DARK


# ── PORTADA ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
r = p.add_run('REPORTE DE ESTADO')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = COLOR_TITLE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('FisioGest — Sistema de Gestión Fisioterapéutica')
r2.font.size = Pt(13)
r2.font.color.rgb = COLOR_GRAY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Fecha: 29 de mayo de 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = COLOR_GRAY

doc.add_paragraph()

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
rl = p_line.add_run('─' * 65)
rl.font.color.rgb = COLOR_TITLE

doc.add_paragraph()

# ── 1. ARQUITECTURA ───────────────────────────────────────────────────
heading1('1. Arquitectura actual')
body('El proyecto combina dos enfoques que no están integrados correctamente:')
bullet('Backend Laravel: rutas en web.php y api.php, vistas Blade en resources/views/')
bullet('Frontend Vue SPA: componentes en resources/components/, router en resources/js/router.js, montado en welcome.blade.php')
body(
    'Esta mezcla provoca conflictos en el enrutamiento, el layout visual y la fuente de datos '
    'que se detallan en las secciones siguientes.'
)

doc.add_page_break()

# ── 2. PROBLEMAS ──────────────────────────────────────────────────────
heading1('2. Problemas detallados')

# BLOQUE A
heading2('Bloque A — Layout / Navegación')

issue_title('A1. Dos sidebars simultáneos')
body(
    'App.vue tiene su propio sidebar con emojis y clases Tailwind CSS. Los componentes '
    'Dashboard.vue, Inventario.vue, etc. usan AppLayout.vue que contiene un segundo sidebar '
    'completo con SVG icons, dropdowns y botón de cerrar sesión. Al navegar a cualquier '
    'página autenticada, ambos se renderizan al mismo tiempo, duplicando la barra lateral '
    'en pantalla.'
)

issue_title('A2. Ruta /citas no existe en el Vue Router')
body(
    'El archivo router.js solo define rutas para: /, /dashboard, /inventario y '
    '/fisioterapeutas. El enlace a "Citas" en AppLayout.vue no lleva a ningún componente '
    'porque no existe Citas.vue. Sin embargo, sí existen lógica en api.php (GET y POST), '
    'vista Blade citas.blade.php y el modal modal-nueva-cita.blade.php — todos sin conectar '
    'al SPA.'
)

issue_title('A3. Link de Pacientes rompe el SPA')
body(
    'En App.vue, el enlace a Pacientes usa <a href="/pacientes"> (enlace HTML nativo) en '
    'lugar de <router-link>. Esto abandona completamente el SPA Vue y carga la vista Blade '
    'de pacientes, perdiendo el estado de la aplicación. En AppLayout.vue, el dropdown de '
    'Pacientes tampoco navega — solo ejecuta un toggle visual.'
)

# BLOQUE B
heading2('Bloque B — Base de datos / Migraciones')

issue_title('B1. Tabla citas tiene fecha_hora (un campo); api.php inserta fecha y hora por separado')
body(
    'La migración 000005 define la columna fecha_hora como un único campo dateTime. '
    'El endpoint POST /api/citas en api.php intenta insertar dos columnas separadas: fecha '
    'y hora, que no existen en la tabla. Esto provoca un error 500 cada vez que se intenta '
    'guardar una cita desde el frontend.'
)

issue_title('B2. ENUM de inventario.estado no coincide con los valores que guarda Vue')
body(
    'La migración define el campo estado como ENUM con valores: disponible, en_uso, '
    'mantenimiento, baja. El componente Inventario.vue guarda los valores: "Available", '
    '"Stock Bajo", "En Uso" (mezcla de inglés y español). SQLite rechazaría esos valores '
    'al intentar insertar un registro.'
)

issue_title('B3. POST de pacientes vía API está incompleto')
body(
    'El endpoint POST /api/pacientes en api.php solo inserta nombre y usuario_id. La tabla '
    'pacientes tiene columnas NOT NULL sin valor por defecto: apellido, fecha_nacimiento, '
    'genero. Cualquier intento de guardar un paciente desde el frontend falla con error 500.'
)

# BLOQUE C
heading2('Bloque C — Controladores / Rutas')

issue_title('C1. InventarioController::index() devuelve una vista Blade, no JSON')
body(
    'El controlador está en la carpeta Api/ y registrado en api.php, pero el método index() '
    'hace return view("inventario", ...) en lugar de return response()->json(...). Cuando '
    'el Dashboard Vue llama a /api/inventario vía axios, recibe HTML en lugar de JSON. '
    'El Promise.allSettled captura el error silenciosamente y los contadores de estadísticas '
    'quedan siempre en cero.'
)

issue_title('C2. InventarioController::store() redirige a un route name inexistente')
body(
    'El método store() redirige a la ruta nombrada "dashboard.inventario.index", pero ese '
    'nombre no está definido en ningún archivo de rutas. El nombre correcto en web.php es '
    '"inventario.index". Al guardar un equipo se produce una excepción de ruta no encontrada.'
)

issue_title('C3. Datos de fisioterapeutas hardcodeados en las rutas')
body(
    'En web.php, las rutas /fisioterapeutas, /pacientes, /citas e /inventario retornan '
    'colecciones PHP estáticas con nombres ficticios (Manrivel Gorado, Barvis Raten, etc.). '
    'La tabla fisioterapeutas existe en la base de datos pero nunca se consulta desde '
    'las rutas web.'
)

# BLOQUE D
heading2('Bloque D — Modelos')

issue_title('D1. Campo marca falta en el $fillable del modelo Inventario')
body(
    'El array $fillable en Inventario.php incluye: nombre, tipo, modelo, estado, cantidad, '
    'descripcion. Omite el campo marca que sí está en la migración y que el formulario Vue '
    'envía como parte del formulario de nuevo equipo. El campo sería silenciosamente '
    'descartado por la protección de mass assignment de Eloquent.'
)

# BLOQUE E
heading2('Bloque E — Autenticación')

issue_title('E1. Login de Vue es hardcodeado, no llama a la API')
body(
    'Login.vue compara directamente las credenciales contra los valores hardcodeados '
    '"admin@fisiogest.com" y "admin123", sin hacer ninguna petición al backend. '
    'El AuthController.php existe con endpoints funcionales de register, login, logout y me, '
    'pero nunca se invoca desde el frontend.'
)

issue_title('E2. Nombres de campos incompatibles entre Vue y AuthController')
body(
    'AuthController espera los campos "correo" y "contraseña" en el request. El formulario '
    'de Login.vue tiene los v-model vinculados a form.email y form.password. Aunque se '
    'conectara el formulario a la API, la validación del controlador fallaría por el '
    'desajuste de nombres.'
)

# BLOQUE F
heading2('Bloque F — Componentes Vue sin conexión a la API')

issue_title('F1. Inventario.vue usa datos completamente hardcodeados')
body(
    'El componente tiene 7 items ficticios definidos directamente en un ref([]). '
    'El servicio inventarioService está definido en api.js y es capaz de llamar a '
    '/api/inventario, pero nunca se importa ni se usa en Inventario.vue. '
    'Los cambios del usuario (agregar/editar) no persisten en la base de datos.'
)

issue_title('F2. No existe Citas.vue')
body(
    'Es la sección más incompleta del frontend. Existe lógica en api.php (rutas GET y POST), '
    'una vista Blade en citas.blade.php y un modal en modal-nueva-cita.blade.php, pero '
    'ningún componente Vue que maneje esta pantalla dentro del SPA.'
)

# BLOQUE G
heading2('Bloque G — Archivos basura')

issue_title("G1. Archivos con nombres de código PHP en la raíz del proyecto")
body(
    "En el directorio FisioGest/FisioGest/ existen archivos cuyos nombres son fragmentos "
    "de código PHP: 'nullable, 'required, validate([, with('success'). Son residuos de "
    "algún copiar/pegar accidental en la terminal o el explorador de archivos. "
    "Deben eliminarse antes de hacer cualquier despliegue."
)

doc.add_page_break()

# ── 3. TABLA DE PRIORIDADES ───────────────────────────────────────────
heading1('3. Tabla de prioridades')

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
for i, txt in enumerate(['Prioridad', 'Bloque', 'Descripción', 'Archivos afectados']):
    hdr_cells[i].text = txt
    set_cell_bg(hdr_cells[i], '074434')
    for para in hdr_cells[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_WHITE

rows_data = [
    ('Alta',  'A1', 'Eliminar sidebar duplicado de App.vue',                   'App.vue'),
    ('Alta',  'A2', 'Crear Citas.vue y agregar ruta al router',                'router.js, AppLayout.vue'),
    ('Alta',  'C1', 'InventarioController devuelve Blade en vez de JSON',      'InventarioController.php'),
    ('Alta',  'B1', 'POST de citas usa columnas inexistentes',                  'api.php'),
    ('Media', 'F1', 'Inventario.vue sin conexión a la API',                    'Inventario.vue, api.js'),
    ('Media', 'B2', 'Valores de estado no coinciden con ENUM de la BD',        'Inventario.vue, migración 000004'),
    ('Media', 'A3', 'Link de Pacientes rompe el SPA',                          'App.vue, AppLayout.vue'),
    ('Media', 'B3', 'POST de pacientes incompleto',                             'api.php'),
    ('Baja',  'D1', 'Campo marca falta en $fillable del modelo',               'Inventario.php (Model)'),
    ('Baja',  'C2', 'Nombre de ruta incorrecto en store()',                    'InventarioController.php'),
    ('Baja',  'E1/E2', 'Login hardcodeado y campos incompatibles con la API',  'Login.vue, AuthController.php'),
    ('Baja',  'G1', 'Archivos basura en directorio raíz',                      'Directorio FisioGest/FisioGest/'),
]

priority_colors = {
    'Alta':  ('FECACA', COLOR_HIGH),
    'Media': ('FEF3C7', COLOR_MED),
    'Baja':  ('DCFCE7', COLOR_LOW),
}

for pri, bloque, desc, archivos in rows_data:
    row = table.add_row().cells
    row[0].text = pri
    row[1].text = bloque
    row[2].text = desc
    row[3].text = archivos

    bg, color = priority_colors[pri]
    set_cell_bg(row[0], bg)

    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

    for para in row[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = color

col_widths = [Cm(2.0), Cm(1.8), Cm(8.8), Cm(5.4)]
for row in table.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w

doc.add_paragraph()

# ── 4. PRÓXIMOS PASOS ─────────────────────────────────────────────────
heading1('4. Próximos pasos recomendados')
body('Se recomienda trabajar en el siguiente orden para desbloquear funcionalidad visible de forma progresiva:')

steps = [
    ('Paso 1', 'Eliminar el sidebar de App.vue y convertirlo en un simple <router-view /> para que AppLayout.vue maneje todo el layout.'),
    ('Paso 2', 'Crear Citas.vue y registrar la ruta /citas en router.js. Conectar el enlace del sidebar.'),
    ('Paso 3', 'Corregir InventarioController::index() para que devuelva JSON. Esto hará que el Dashboard muestre datos reales.'),
    ('Paso 4', 'Corregir el endpoint POST /api/citas para usar el campo fecha_hora combinado.'),
    ('Paso 5', 'Conectar Inventario.vue a la API (reemplazar los 7 items hardcodeados por llamadas reales).'),
    ('Paso 6', 'Alinear los valores de estado de inventario con el ENUM definido en la base de datos.'),
    ('Paso 7', 'Completar el POST de pacientes con todos los campos requeridos por la migración.'),
    ('Paso 8', 'Corregir detalles menores: $fillable del modelo, nombre de ruta en store(), archivos basura.'),
]

for num, text in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(num + ': ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = COLOR_TITLE
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = COLOR_BODY

# ── GUARDAR ───────────────────────────────────────────────────────────
output = (
    r'c:\Users\abeli\Downloads\Clases ing. Sistemas\Ciclo 5'
    r'\Programacion computacional 4\Bloque 3\Actividades'
    r'\Proyectos VS Code\Proyecto final de la materia'
    r'\FisioGest\Reporte_Estado_FisioGest.docx'
)
doc.save(output)
print('Guardado en:', output)
